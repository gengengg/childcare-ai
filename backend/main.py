import io
import json
import os
import re
from pathlib import Path
from typing import List, Optional
from urllib.parse import quote

import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from lxml import etree
from openai import OpenAI
from pydantic import BaseModel

load_dotenv(dotenv_path=Path(__file__).parent / ".env")

api_key = os.getenv("OPENAI_API_KEY")
model_name = os.getenv("OPENAI_MODEL", "gpt-4o")

client = OpenAI(api_key=api_key, timeout=25.0, max_retries=0)
client_observation = OpenAI(api_key=api_key, timeout=90.0, max_retries=0)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ──────────────── 요청 모델 ────────────────

class ActivityItem(BaseModel):
    category: str
    title: str
    memo: Optional[str] = ""


class DailyRecordRequest(BaseModel):
    child_name: str
    class_name: str
    date: str
    activities: List[ActivityItem]
    meal_note: Optional[str] = ""
    nap_note: Optional[str] = ""
    health_note: Optional[str] = ""
    images: Optional[List[str]] = []
    style_guide: Optional[str] = ""
    emoji_enabled: Optional[bool] = True


class StyleSampleItem(BaseModel):
    kind: str  # "text" 또는 "image"
    text: Optional[str] = ""
    image_base64: Optional[str] = ""


class AnalyzeStyleRequest(BaseModel):
    samples: List[StyleSampleItem]


class ObservationActivityItem(BaseModel):
    category: str
    title: str
    memo: Optional[str] = ""


class ObservationSourceRecord(BaseModel):
    date: str
    teacher_final: str
    activities: Optional[List[ObservationActivityItem]] = []
    meal_note: Optional[str] = ""
    nap_note: Optional[str] = ""
    health_note: Optional[str] = ""


class ObservationRequest(BaseModel):
    child_name: str
    class_name: str
    child_age: Optional[int] = 2
    start_date: str
    end_date: str
    child_observation_notes: Optional[str] = ""
    records: List[ObservationSourceRecord]


class ObservationCategories(BaseModel):
    기본생활습관: str
    신체건강: str
    의사소통: str
    사회관계: str
    예술경험: str
    자연탐구: str
    총평: str


class ExportDocxRequest(BaseModel):
    child_name: str
    class_name: str
    observation_date: str
    teacher_name: str
    semester: int
    content: ObservationCategories


# ──────────────── 공통 유틸 ────────────────

def call_openai_text(prompt: str, instructions: str, max_output_tokens: int = 1800) -> str:
    response = client.responses.create(
        model=model_name,
        instructions=instructions,
        input=prompt,
        max_output_tokens=max_output_tokens,
    )
    return response.output_text


def _detect_mime(b64: str) -> str:
    if b64.startswith("/9j/"):
        return "image/jpeg"
    if b64.startswith("iVBORw0KGgo"):
        return "image/png"
    if b64.startswith("UklGR"):
        return "image/webp"
    if b64.startswith("R0lGOD"):
        return "image/gif"
    return "image/jpeg"


def call_openai_with_images(prompt: str, instructions: str, images: List[str], max_output_tokens: int = 1800) -> str:
    if not images:
        return call_openai_text(prompt, instructions, max_output_tokens)

    content = []
    for b64 in images:
        mime = _detect_mime(b64)
        content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:{mime};base64,{b64}",
                "detail": "high",
            },
        })
    content.append({"type": "text", "text": prompt})

    response = client_observation.chat.completions.create(
        model=model_name,
        messages=[
            {"role": "system", "content": instructions},
            {"role": "user", "content": content},
        ],
        max_tokens=max_output_tokens,
    )
    return response.choices[0].message.content


def make_fallback_daily_record(request: DailyRecordRequest) -> str:
    lines = []
    for i, act in enumerate(request.activities, 1):
        lines.append(
            f"{i}. [{act.category}] 활동명: {act.title}\n"
            f"학부모님, 오늘은 {act.title} 활동으로 즐거운 시간을 보냈습니다. "
            f"{request.child_name}이(가) 열심히 참여하는 모습이 무척 대견했습니다."
        )

    daily_notes = []
    if (request.meal_note or "").strip():
        daily_notes.append(f"식사: {request.meal_note.strip()}")
    if (request.nap_note or "").strip():
        daily_notes.append(f"수면: {request.nap_note.strip()}")
    if (request.health_note or "").strip():
        daily_notes.append(f"특이사항: {request.health_note.strip()}")

    result = "\n\n".join(lines)
    if daily_notes:
        result += "\n\n" + " / ".join(daily_notes)
    result += "\n\n행복한 하루 보내세요."
    return result


# ──────────────── docx 생성 유틸 ────────────────

def _set_row_min_height(row, height_emu: int):
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = etree.Element(qn("w:trPr"))
        tr.insert(0, trPr)
    trHeight = trPr.find(qn("w:trHeight"))
    if trHeight is None:
        trHeight = etree.SubElement(trPr, qn("w:trHeight"))
    trHeight.set(qn("w:val"), str(int(height_emu / 635)))
    trHeight.set(qn("w:hRule"), "atLeast")


def _set_cell(cell, text: str, bold=False, center=False, size=11, font="맑은 고딕"):
    para = cell.paragraphs[0]
    para.clear()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), font)


def build_observation_docx(req: ExportDocxRequest) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(6)
    title_run = title_para.add_run(f"{req.semester}학기  영·유아 관찰평가")
    title_run.font.name = "맑은 고딕"
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    rPr = title_run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    table = doc.add_table(rows=10, cols=4)
    table.style = "Table Grid"
    col_widths = [Cm(3.0), Cm(4.83), Cm(4.83), Cm(4.84)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]

    r0 = table.rows[0]
    _set_cell(r0.cells[0], "반  명", bold=True, center=True)
    _set_cell(r0.cells[1], req.class_name, center=True)
    _set_cell(r0.cells[2], "관찰일시", bold=True, center=True)
    _set_cell(r0.cells[3], req.observation_date, center=True)

    r1 = table.rows[1]
    _set_cell(r1.cells[0], "관찰교사", bold=True, center=True)
    merged = r1.cells[1].merge(r1.cells[2]).merge(r1.cells[3])
    _set_cell(merged, f"{req.teacher_name}                               (서명또는 인)")

    r2 = table.rows[2]
    _set_cell(r2.cells[0], "생활영역", bold=True, center=True)
    merged2 = r2.cells[1].merge(r2.cells[2]).merge(r2.cells[3])
    _set_cell(merged2, "관   찰   내   용", bold=True, center=True)

    areas = [
        ("기본생활습관", req.content.기본생활습관),
        ("신체·건강",   req.content.신체건강),
        ("의사소통",    req.content.의사소통),
        ("사회관계",    req.content.사회관계),
        ("예술경험",    req.content.예술경험),
        ("자연탐구",    req.content.자연탐구),
        ("총   평",    req.content.총평),
    ]
    for i, (label, text) in enumerate(areas):
        row = table.rows[3 + i]
        _set_cell(row.cells[0], label, bold=True, center=True, size=10)
        merged = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
        _set_cell(merged, text, size=10)

    _set_row_min_height(table.rows[0], Cm(0.9))
    _set_row_min_height(table.rows[1], Cm(0.9))
    _set_row_min_height(table.rows[2], Cm(0.7))
    for i in range(3, 9):
        _set_row_min_height(table.rows[i], Cm(3.1))
    _set_row_min_height(table.rows[9], Cm(3.6))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────── 엔드포인트 ────────────────

@app.get("/")
def health_check():
    return {"status": "ok", "message": "childcare-ai backend is running"}


@app.post("/generate-daily-record")
def generate_daily_record(request: DailyRecordRequest):
    if not api_key:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY가 설정되지 않았습니다.")

    has_images = bool(request.images)

    # 허용 영역 목록 (표준 보육과정 5영역, 프론트 ACTIVITY_CATEGORIES와 동일)
    VALID_AREAS = {'신체운동·건강', '의사소통', '사회관계', '예술경험', '자연탐구'}

    activities_text = ""
    for i, act in enumerate(request.activities, 1):
        memo = act.memo.strip() if act.memo else "없음"
        cat = act.category.strip() if act.category else ""
        if cat in VALID_AREAS:
            area_line = f"출력 영역(고정): [{cat}]  ← 이 값을 그대로 사용"
        else:
            area_line = (
                "출력 영역(AI 선택 필요): 아래 표준 보육과정 5영역 중 활동명에 맞는 것을 반드시 하나 선택\n"
                "  → 신체운동·건강 / 의사소통 / 사회관계 / 예술경험 / 자연탐구\n"
                "  → 이 목록 외의 어떤 표현도 사용 불가 ([기본생활] [언어·의사소통] [창의·탐구] [놀이] [체험활동] [실내활동] 등 모두 금지)"
            )
        activities_text += f"\n[활동{i}]\n  {area_line}\n  활동명: {act.title}\n  교사 메모: {memo}\n"

    daily_notes = []
    if (request.meal_note or "").strip():
        daily_notes.append(f"식사: {request.meal_note.strip()}")
    if (request.nap_note or "").strip():
        daily_notes.append(f"수면: {request.nap_note.strip()}")
    if (request.health_note or "").strip():
        daily_notes.append(f"특이사항: {request.health_note.strip()}")
    daily_notes_text = ", ".join(daily_notes) if daily_notes else "없음"

    photo_section = ""
    if has_images:
        photo_section = """
[사진 활용 — 최우선 규칙]
첨부된 사진을 먼저 아주 자세히 살펴봐. 마치 학부모가 그 자리에서 아이를 직접 보는 듯한 생생함을 만드는 게 목표야.
- 사진 속 아이의 실제 행동, 손동작, 시선, 표정, 자세, 주변 사물·교구·색깔까지 구체적으로 묘사해.
  예) "싸인펜을 손에 꼭 쥐고", "고개를 갸웃하며 진지하게 들여다보고", "눈을 동그랗게 뜨고", "입을 오물오물하며", "두 손으로 조심스럽게 받쳐 들고"
- 사진에 보이는 교구·재료·색깔·모양을 그대로 언급해 (예: 파란색 바다 매트, 알록달록 물고기 모양 조각, 하얀색 종이, 잠자리채 등).
- 사진에 없는 내용·인물·행동을 절대 만들어내지 마. 추측되는 감정은 사진 속 표정에서 명확히 보일 때만 자연스럽게 한 줄로 덧붙여.
- 활동명·교사 메모는 큰 틀만 참고하고, 구체적인 장면 묘사는 반드시 사진에서 가져와.
"""

    style_guide_text = (request.style_guide or "").strip()
    if style_guide_text:
        style_section = f"""
[사용자 알림장 문체 — 최우선 적용]
이 교사는 평소 자신만의 알림장 문체를 가지고 있어. 아래에 정리된 문체 가이드를 알림장 전체에 자연스럽게 반영해야 해.
이 가이드는 아래의 '문체 규칙' 기본 예시보다 우선해. 단, 표준 보육과정 5영역 규정과 [영역명] 소제목 출력 형식은 절대 바꾸지 마.

{style_guide_text}

위 문체 가이드의 어휘·문장 길이·시작 표현·맺음 표현·즐겨 쓰는 표현·말투(존댓말 정도)·특수문자 사용 여부 등을 그대로 따라가. 단, 학부모를 무시하거나 무례하게 들리는 표현은 사용하지 마.
(이모지 사용 여부는 아래 '이모지 사용 규칙'을 절대적으로 우선해서 따라. 문체 가이드의 이모지 관련 내용은 무시할 것.)
"""
        tone_instructions = (
            "사용자가 기존에 쓰던 알림장 문체 가이드가 입력되어 있다면, 그 문체를 충실히 따라 작성해."
            "기본 문체 규칙(~어요체 80%)은 그 문체 가이드와 충돌할 때 가이드를 우선해."
        )
    else:
        style_section = ""
        tone_instructions = (
            "~어요 체 위주(80% 이상)로 작성하고, ~하십시오·~합니다는 금지하는 기본 문체 규칙을 따라."
        )

    if request.emoji_enabled is False:
        emoji_section = """=== 이모지(이모티콘) 사용 규칙 — 반드시 준수 ===
- 이번 알림장에는 이모지(이모티콘)를 절대 사용하지 마.
- 어떤 영역, 어떤 문장에도 이모지·그림문자·심볼(🌟 🎨 😊 ✨ 등)을 포함하지 마.
- 문체 가이드에 이모지가 있어도 이 규칙이 우선이다. 텍스트로만 따뜻한 묘사를 만들어.
- 의성어·의태어·구체적 행동 묘사로 생생함을 살려 (이모지 없이도 충분히 풍부하게)."""
        emoji_instruction = "이번 알림장은 이모지를 절대 사용하지 마. 텍스트로만 작성해."
    else:
        emoji_section = """=== 이모지(이모티콘) 사용 규칙 ===
- 활동 하나당 본문에 2~3개의 이모지를 자연스럽게 섞어 써. (반드시 2개 이상, 최대 3개. 0개나 4개 이상 금지)
- 이모지는 활동 성격에 맞는 것으로 골라. 영역별 추천 예시:
  · 신체운동·건강: 🤸 🏃 ⚽ 🌟 💪 🦋 🌈
  · 의사소통·언어: 📚 🗣️ ✏️ 💬 📖
  · 사회관계·친구: 🤗 👫 💖 ✨ 🥰
  · 예술경험·미술·음악: 🎨 🖍️ 🎵 🎶 🌸 🍀 🎭
  · 자연탐구·관찰: 🔍 🌱 🐟 🐞 🌊 ☀️ 🍂
  · 식사·간식: 🍚 🥄 🍎 😋
  · 수면·휴식: 😴 💤 🛏️
- 이모지는 문장 사이사이에 자연스럽게 배치해. 한 곳에 몰아서 쓰지 마.
- 이모지는 문장 끝에만 몰지 말고, 관련된 단어 바로 뒤(중간)에 붙여도 좋다. 단어와 이모지 사이에는 공백을 넣지 않는다.
  · 좋은 예: "물놀이🌊를 하며 까르르 웃었어요."
  · 좋은 예: "블록🧱을 높이 쌓아 올리며 뿌듯한 표정을 지었어요."
  · 좋은 예: "퐁당퐁당🌊 물놀이를 하며 까르르 웃었어요😊"
  · 피할 예 (끝에만 몰림): "오늘은 물놀이를 하며 까르르 웃었어요 🌊😊✨"
- 이모지가 특정 명사(교구/자연물/음식 등)를 직접 가리킬 때는 그 단어 뒤에 붙이는 걸 우선한다.
  감정·분위기를 나타낼 때는 문장 끝이나 쉼표 앞에 배치해도 자연스럽다.
- 소제목(영역명 옆 제목)에는 이모지를 쓰지 마. 본문에만 사용해.
- 같은 이모지를 한 활동 안에서 반복하지 마."""
        emoji_instruction = (
            "활동 본문 하나당 이모지를 반드시 2~3개(2개 이상, 3개 이하) 자연스럽게 배치해. "
            "영역 성격에 맞는 이모지를 골라 쓰고, 한 곳에 몰지 말고 문장 사이사이에 흩어 배치해. "
            "명사(교구·자연물·음식 등)를 가리키는 이모지는 그 단어 바로 뒤에 공백 없이 붙여 쓰는 것도 자연스럽다 "
            "(예: '물놀이🌊를 했어요'). 문장 끝에만 몰아 붙이지 마. 소제목에는 이모지를 쓰지 마."
        )

    prompt = f"""
너는 어린이집 보육교사의 알림장 작성을 돕는 AI야.
{photo_section}{style_section}
아래 활동 정보를 바탕으로 학부모님께 전달할 따뜻한 알림장을 작성해줘.

=== 출력 형식 (반드시 준수) ===
숫자. [영역명] 소제목
내용 (250~350자, ~어요 체 위주)

예시:
1. [신체운동·건강] 씩씩하게 뛰어놀아요
오늘 한은이는 ...

※ 영역명은 반드시 대괄호 [ ] 안에 넣어야 해.

활동이 여러 개면 성격에 따라 2개 묶음으로 나눠 작성.
활동이 1개면 1개만 작성.

=== 이름 규칙 ===
아이 이름에서 첫 글자(성)를 제거하고 이름만 사용.
  정한은 → 한은 / 김채린 → 채린 / 나연서 → 연서
이름 끝에 받침 있으면 '이'를 붙인 뒤 조사 연결.
  한은(받침ㄴ) → 한은이가, 한은이는, 한은이와
  채린(받침ㄴ) → 채린이가, 채린이는
이름 끝에 받침 없으면 바로 조사 연결.
  연서 → 연서가, 연서는
성+이름 전체 사용 금지 (예: 정한은이~, 정한은 친구들과~ 모두 금지)

=== 문체 규칙 (사용자 문체 가이드가 위에 있으면 그것을 우선) ===
~어요 / ~했어요 / ~즐겼어요 위주로 작성 (전체 문장의 80% 이상)
~습니다는 최소한으로만 (20% 이하)
~하십시오 / ~하시기 바랍니다 / ~합니다 절대 금지
선생님이 학부모님께 아이 하루를 다정하게 이야기해주는 느낌으로 작성

=== 생생한 묘사 규칙 (매우 중요) ===
"성의 없어 보인다"는 피드백을 받지 않도록, 학부모가 그 장면을 눈앞에서 보는 듯한 알림장을 써야 해.

1) 유아교육 의성어·의태어를 본문 안에 자연스럽게 녹여 적극적으로 사용해. 활동 성격에 맞는 것을 골라 써.
   - 물·바다 놀이: 퐁당, 첨벙첨벙, 출렁출렁, 쏙쏙
   - 미술·쓰기: 끼적끼적, 쓱싹쓱싹, 알록달록, 동글동글, 쭉쭉
   - 신체·움직임: 폴짝폴짝, 살금살금, 뒤뚱뒤뚱, 깡충깡충, 사뿐사뿐
   - 음식·먹기: 냠냠, 오물오물, 호로록, 꿀꺽
   - 탐색·관찰: 빤히, 갸웃갸웃, 말똥말똥, 두리번두리번
   - 감정·표정: 방긋, 싱긋, 까르르, 헤헤
   (위는 예시일 뿐 — 사진과 활동 맥락에 맞는 것을 골라 써. 억지로 끼워 넣지는 마.)

2) 유아교육 현장 용어를 자연스럽게 활용해.
   - "끼적이기" (낙서·그리기를 표현하는 영유아 보육 용어)
   - "탐색해 보고", "탐색 활동", "오감 놀이", "조작 활동"
   - "온몸으로 헤치며", "직접 매칭해 보며", "쏙쏙 건져 보며"
   - "○○ 세상 속으로 빠져보았습니다", "○○ 활동을 해보았어요"

3) 구체적인 행동·표정·감각을 한 문장에 하나씩은 넣어. 추상적인 "즐거워했어요", "재미있어했어요"만 반복하지 마.
   나쁜 예) "오늘 친구들과 즐겁게 활동했어요. 잘 참여했습니다."
   좋은 예) "싸인펜을 손에 꼭 쥐고 활동지를 진지하게 들여다보며 멋지게 끼적이기 활동을 해보았어요. 스스로 해내려는 의지를 보이며 집중하는 모습이 무척 대견했답니다."

4) 색깔·모양·재료·교구 이름을 구체적으로 적어. ("파란색 종이 바다", "하얀색 물고기 모양 조각", "알록달록 동글동글한", "잠자리채(뜰채)")

5) 영역별 소제목은 활동 핵심을 살린 짧고 정겨운 문장체로 지어줘.
   예) "시원한 바다에서 놀이해요", "내가 혼자 할 수 있어요", "씩씩하게 뛰어놀아요"

{emoji_section}

=== 기타 규칙 ===
마크다운 기호(** ## * 등) 사용 금지
이미지 파일명 포함 금지
없는 사실 금지
"학부모님,"으로 시작 금지
식사/수면/특이사항 메모가 있으면 마지막에 한 줄로 자연스럽게 덧붙이고 짧은 인사로 마무리 (건강 외에도 컨디션·기분·친구 관계·전달 사항 등 어떤 내용이든 들어올 수 있음)

=== 입력 정보 ===
아이 이름: {request.child_name}
반: {request.class_name}
날짜: {request.date}

오늘의 활동:
{activities_text}
일상 메모: {daily_notes_text}

알림장만 작성해줘.
"""

    instructions = (
        "너는 어린이집 보육교사의 알림장 작성을 돕는 전문 AI야. "
        "활동 영역은 입력에 '출력 영역(고정)'으로 표시된 경우 그대로 쓰고, "
        "'출력 영역(AI 선택 필요)'인 경우 표준 보육과정 5영역 — 신체운동·건강 / 의사소통 / 사회관계 / 예술경험 / 자연탐구 — 중 하나만 선택해. "
        "[기본생활]·[언어·의사소통]·[창의·탐구]·[놀이]·[실내활동]·[체험활동] 등 목록 외 영역명은 오답이야. "
        "사진이 있으면 사진 속 장면(아이의 손동작·표정·시선·교구·색깔)을 최우선으로 구체적으로 묘사해. "
        "유아교육 의성어·의태어(퐁당, 끼적끼적, 쏙쏙, 폴짝폴짝, 알록달록, 갸웃갸웃 등)와 현장 용어('끼적이기', '탐색해 보고', '온몸으로 헤치며')를 활동 맥락에 맞게 자연스럽게 녹여 써. "
        "추상적인 '즐거웠어요/재미있었어요'만 반복하지 말고, 구체적 행동·표정·감각을 살려 학부모가 그 장면을 눈앞에서 보는 듯하게 작성해. "
        f"{emoji_instruction} "
        "아이 이름은 성(첫 글자) 제거 후 이름만 쓰고, 받침 있으면 '이'를 붙여(한은→한은이가). "
        f"{tone_instructions} "
        "마크다운 없이 출력해."
    )

    try:
        draft = call_openai_with_images(
            prompt=prompt,
            instructions=instructions,
            images=request.images or [],
            max_output_tokens=1200,
        )
        return {"draft": draft, "source": "openai"}
    except Exception as error:
        print("OPENAI DAILY RECORD ERROR:", repr(error))
        return {
            "draft": make_fallback_daily_record(request),
            "source": "fallback",
            "warning": f"OpenAI 호출 실패: {repr(error)}",
        }


@app.post("/analyze-style")
def analyze_style(request: AnalyzeStyleRequest):
    if not api_key:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY가 설정되지 않았습니다.")

    samples = [s for s in (request.samples or []) if (s.text or "").strip() or (s.image_base64 or "").strip()]
    if not samples:
        raise HTTPException(status_code=400, detail="분석할 알림장 샘플이 없습니다.")
    if len(samples) > 5:
        raise HTTPException(status_code=400, detail="알림장 샘플은 최대 5개까지 분석할 수 있습니다.")

    texts: List[str] = []
    images: List[str] = []
    for idx, s in enumerate(samples, 1):
        if s.kind == "text" and (s.text or "").strip():
            texts.append(f"[샘플 {idx} — 텍스트]\n{s.text.strip()}")
        elif s.kind == "image" and (s.image_base64 or "").strip():
            texts.append(f"[샘플 {idx} — 이미지 첨부됨: 이미지 속 알림장 본문을 읽어서 함께 반영]")
            images.append(s.image_base64)

    samples_text = "\n\n".join(texts) if texts else "(텍스트 샘플 없음, 이미지만 제공됨)"

    prompt = f"""
아래는 어린이집 보육교사 한 명이 기존에 쓰던 알림장 샘플들이야 (텍스트 또는 이미지).
이 교사가 앞으로 새 알림장을 쓸 때 동일한 문체로 자동 작성될 수 있도록,
샘플들에서 공통적으로 나타나는 문체적 특징을 한국어로 정리해줘.

분석해야 할 항목:
1) 문장 종결 어미 패턴 (예: ~어요, ~했답니다, ~네요, ~입니다 등 — 비율도 추정)
2) 평균 문장 길이와 한 알림장의 평균 길이(글자 수 범위)
3) 자주 등장하는 시작 표현 / 마무리 인사 표현
4) 자주 쓰는 어휘·표현·의성어·의태어 (구체적인 단어 목록)
5) 호칭 (학부모님, 어머님, 부모님 등) 및 아이 지칭 방법
6) 이모지·특수문자·줄바꿈 패턴
7) 활동을 설명하는 톤 (담백 / 다정 / 격식 / 친근 등)
8) 그 외 이 교사만의 두드러진 습관

반드시 아래 형식의 plain text로 응답해 (마크다운, JSON, 코드블록 금지):

문장 종결: ...
평균 길이: ...
시작 표현: ...
마무리 표현: ...
자주 쓰는 어휘: ...
호칭/아이 지칭: ...
이모지/특수문자: ...
전체적인 톤: ...
기타 특징: ...

샘플:
{samples_text}
"""

    instructions = (
        "너는 한국어 보육 알림장 문체 분석가야. "
        "주어진 샘플들에서 공통된 문체적 특징을 정확하고 간결하게 한국어로 정리해. "
        "샘플에 없는 추측은 하지 말고, 마크다운/코드블록 없이 plain text로만 응답해."
    )

    try:
        if images:
            guide = call_openai_with_images(
                prompt=prompt,
                instructions=instructions,
                images=images,
                max_output_tokens=900,
            )
        else:
            guide = call_openai_text(prompt, instructions, max_output_tokens=900)
        return {"style_guide": (guide or "").strip()}
    except Exception as error:
        print("STYLE ANALYSIS ERROR:", repr(error))
        raise HTTPException(status_code=500, detail=f"문체 분석 중 오류: {repr(error)}")


@app.post("/generate-observation-structured")
def generate_observation_structured(request: ObservationRequest):
    if not api_key:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY가 설정되지 않았습니다.")
    if len(request.records) == 0:
        raise HTTPException(status_code=400, detail="보육일지 기록이 없습니다.")
    return _run_structured_observation(request)


@app.post("/export-observation-docx")
def export_observation_docx(request: ExportDocxRequest):
    try:
        docx_bytes = build_observation_docx(request)
        filename = f"관찰일지_{request.child_name}_{request.observation_date}.docx"
        encoded_filename = quote(filename, safe="")
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
        )
    except Exception as error:
        print("DOCX EXPORT ERROR:", repr(error))
        raise HTTPException(status_code=500, detail=f"docx 생성 중 오류: {repr(error)}")


# ──────────────── 내부 함수 ────────────────

def _extract_json(text: str) -> str:
    fence = re.search(r'```(?:json)?\s*([\s\S]+?)\s*```', text)
    if fence:
        return fence.group(1)
    try:
        start = text.index('{')
        end = text.rindex('}') + 1
        return text[start:end]
    except ValueError:
        return text.strip()


def _fallback_observation_categories() -> dict:
    placeholder = "AI 생성에 실패했습니다. 직접 내용을 입력해주세요."
    return {k: placeholder for k in ["기본생활습관", "신체건강", "의사소통", "사회관계", "예술경험", "자연탐구", "총평"]}


_AGE_GUIDANCE = {
    0: "감각 탐색과 신체 감각 인식, 양육자와의 애착 형성, 옹알이와 기초 의사소통 시작 단계.",
    1: "혼자 걷기·뛰기 등 대근육 발달, 초기 언어(단어) 사용 시작, 대상영속성 이해.",
    2: "언어 폭발기(두 단어 조합→짧은 문장), 자아 표현과 자기주장 강해짐. 병행 놀이, 모방 놀이 활발.",
    3: "상상·역할 놀이 발달, 또래와의 상호작용 시작, 언어 급성장(질문 많음).",
    4: "협동 놀이와 규칙 이해·준수, 감정 조절 능력 발달, 복잡한 언어 표현.",
    5: "문해력·수 개념 기초 형성, 복잡한 사회적 상호작용, 계획·목표 지향 행동.",
}


def _run_structured_observation(request: ObservationRequest) -> dict:
    age = request.child_age if request.child_age is not None else 2
    age_guidance = _AGE_GUIDANCE.get(age, _AGE_GUIDANCE[2])

    records_text = ""
    for r in request.records:
        activities_summary = ""
        if r.activities:
            for a in r.activities:
                activities_summary += f"  - [{a.category}] {a.title}"
                if a.memo:
                    activities_summary += f": {a.memo}"
                activities_summary += "\n"
        else:
            activities_summary = "  (활동 정보 없음)\n"

        records_text += f"""
날짜: {r.date}
활동:
{activities_summary}최종 알림장:
{r.teacher_final}
식사: {r.meal_note or '없음'} / 수면: {r.nap_note or '없음'} / 특이사항: {r.health_note or '없음'}
---
"""

    child_notes_section = ""
    if (request.child_observation_notes or "").strip():
        child_notes_section = f"""
교사 관찰 메모 (이 아이에 대해 교사가 직접 기록한 내용):
{request.child_observation_notes}
---
"""

    prompt = f"""
너는 어린이집 보육교사의 관찰일지 작성을 돕는 AI야.
아래 알림장 기록과 교사 관찰 메모를 바탕으로 7개 생활영역 관찰일지를 작성해줘.

아이 발달 단계: 만 {age}세
만 {age}세 발달 특성: {age_guidance}

작성 원칙:
- 반드시 제공된 기록과 관찰 메모에 근거해서만 작성한다.
- 기록에 없는 내용은 절대 만들지 않는다.
- 만 {age}세 발달 수준에 맞게 작성한다.
- 각 영역 80~120자, 총평 100~150자로 간결하게 작성한다.
- 실제 어린이집 관찰일지 문체로 자연스럽게 서술한다.

아이 이름: {request.child_name}
반: {request.class_name}
기간: {request.start_date} ~ {request.end_date}
{child_notes_section}
알림장 기록:
{records_text}

반드시 아래 JSON 형식으로만 응답해:
{{
  "기본생활습관": "...",
  "신체건강": "...",
  "의사소통": "...",
  "사회관계": "...",
  "예술경험": "...",
  "자연탐구": "...",
  "총평": "..."
}}
"""

    try:
        response = client_observation.responses.create(
            model=model_name,
            instructions=(
                f"너는 어린이집 만 {age}세반 보육교사의 관찰일지 작성을 돕는 AI야. "
                "제공된 기록에 근거해서만 작성하고, 반드시 JSON 형식으로만 응답해."
            ),
            input=prompt,
            max_output_tokens=1800,
        )
        raw = _extract_json(response.output_text)
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print("JSON ERROR:", repr(e))
        return _fallback_observation_categories()
    except Exception as e:
        print("OBSERVATION ERROR:", repr(e))
        return _fallback_observation_categories()


# ──────────────── 월간계획안 파싱 ────────────────

def _clean_cell(v) -> str:
    return (v or "").strip().replace("\n", " ").replace("\r", " ")


def _extract_pdf_text(content: bytes) -> str:
    """
    PDF의 표를 [행 | 셀1 | 셀2 | ...] 형태로 보존.
    월간계획안은 '열 = 주' 구조라서 표 구조를 살려야 AI가 주 배정을 올바르게 함.
    """
    parts: List[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page_no, page in enumerate(pdf.pages, 1):
            tables = page.extract_tables() or []
            for tbl_i, tbl in enumerate(tables, 1):
                parts.append(f"\n=== 표 {page_no}-{tbl_i} 시작 ===")
                for row_i, row in enumerate(tbl, 1):
                    cells = [_clean_cell(c) for c in row]
                    parts.append(f"행{row_i}: | " + " | ".join(cells) + " |")
                parts.append(f"=== 표 {page_no}-{tbl_i} 끝 ===\n")
            # 표 밖의 서두 텍스트(주제, 교사 기대 등) 별도 섹션
            txt = page.extract_text() or ""
            if txt.strip():
                parts.append(f"\n--- 페이지 {page_no} 본문 ---")
                parts.append(txt)
    return "\n".join(parts)


def _extract_docx_text(content: bytes) -> str:
    """
    DOCX의 표를 행/열 구조로 보존. 문단은 별도 섹션.
    """
    doc = Document(io.BytesIO(content))
    parts: List[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for tbl_i, table in enumerate(doc.tables, 1):
        parts.append(f"\n=== 표 {tbl_i} 시작 ===")
        for row_i, row in enumerate(table.rows, 1):
            cells = [_clean_cell(c.text) for c in row.cells]
            parts.append(f"행{row_i}: | " + " | ".join(cells) + " |")
        parts.append(f"=== 표 {tbl_i} 끝 ===\n")
    return "\n".join(parts)


@app.post("/parse-monthly-plan")
async def parse_monthly_plan(file: UploadFile = File(...)):
    if not api_key:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY가 설정되지 않았습니다.")

    filename = (file.filename or "").lower()
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="빈 파일입니다.")
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="10MB 이하 파일만 업로드할 수 있어요.")

    if filename.endswith(".pdf"):
        try:
            raw_text = _extract_pdf_text(content)
        except Exception as e:
            print("PDF PARSE ERROR:", repr(e))
            raise HTTPException(status_code=400, detail="PDF에서 텍스트를 추출하지 못했어요.")
    elif filename.endswith(".docx"):
        try:
            raw_text = _extract_docx_text(content)
        except Exception as e:
            print("DOCX PARSE ERROR:", repr(e))
            raise HTTPException(status_code=400, detail="DOCX에서 텍스트를 추출하지 못했어요.")
    else:
        raise HTTPException(
            status_code=400,
            detail="PDF 또는 DOCX 파일만 지원해요. HWP는 한글에서 PDF로 변환 후 업로드해 주세요.",
        )

    if not raw_text.strip():
        raise HTTPException(status_code=400, detail="파일에서 텍스트를 찾을 수 없어요.")

    # 너무 길면 앞부분만 (월간계획안은 보통 2~3페이지)
    if len(raw_text) > 20000:
        raw_text = raw_text[:20000]

    prompt = f"""
아래는 어린이집 월간 놀이 계획안 문서의 텍스트야.
문서는 "=== 표 X 시작 === ... === 표 X 끝 ===" 로 감싸진 표 섹션과, "--- 페이지 N 본문 ---" 으로 표시된 서두 텍스트로 구성되어 있어.

【1단계: 표의 방향을 먼저 판단】
어린이집마다 표 구조가 두 가지 형태가 있어. 표 안에서 "1주 / 2주 / 3주 / 4주" 라는 헤더가 **어디에 나오는지**로 판단해:

▶ 형태 A (열 = 주): 헤더 행 하나에 [빈칸, 1주, 2주, 3주, 4주] 가 가로로 나열됨.
   → **같은 열(같은 위치의 셀)에 있는 값은 모두 같은 주.**
   → 각 데이터 행의 첫 번째 열은 카테고리 이름(놀이, 바깥놀이, 등원, 안전교육 등)이므로 무시.
   → 두 번째 셀 = 1주 활동, 세 번째 셀 = 2주 활동, ... 이런 식으로 열 번호로 판단.

▶ 형태 B (행 = 주): 각 행의 첫 번째 열에 "1주", "2주", "3주", "4주" 가 세로로 나열됨.
   → **같은 행에 있는 모든 셀은 같은 주.**
   → 그 행의 두 번째 셀부터의 모든 값이 해당 주의 활동.
   → 헤더가 [주, 카테고리1, 카테고리2, ...] 형태라면 카테고리 이름은 무시.

▶ 어느 쪽인지 모호하면 헤더 위치와 셀 개수 패턴으로 판별하고, 판별한 형태를 일관되게 적용.

【2단계: 활동 추출 규칙】
- 셀 안에 여러 활동이 "- 활동A - 활동B" 처럼 나열되어 있으면 각각을 분리해서 배열에 넣어.
- 앞의 "- " 나 불릿 기호는 제거.
- 카테고리 접두 태그 "(바깥놀이)", "(실내대체)", "(성폭력)", "(생활동요)" 등은 활동명에 유지.
- 빈 셀("", "-", 공백만)이나 카테고리 이름 셀은 무시.
- **순서 보존: 원본 문서에 나오는 순서 그대로 activities 배열에 넣어야 해.**
  - 형태 A(열=주)일 때: 표의 위쪽 행부터 아래쪽 행 순서대로. 한 셀 안에 여러 활동이 있으면 그 셀 안의 순서 그대로.
  - 형태 B(행=주)일 때: 그 행의 왼쪽 셀부터 오른쪽 셀 순서대로. 한 셀 안에 여러 활동이 있으면 셀 안의 순서 그대로.
  - 절대 알파벳/가나다순으로 재정렬하거나, 카테고리별로 묶어서 순서를 바꾸지 마.

【2-1단계: 활동이 아닌 카테고리는 완전 제외】
아래 카테고리의 행/열에 있는 내용은 **매일 반복되는 일상 흐름**이라 활동으로 보지 말고 activities 배열에서 **완전히 제외**:
- 등원 및 맞이하기 / 등원 / 하원 / 귀가 / 통합보육
- 일상생활 / 일상 생활 (점심, 간식, 손씻기, 낮잠, 휴식, 배변, 기저귀갈이, 이닦기 등)
- 식사 / 급식 / 오전간식 / 오후간식
- 정리정돈 / 정리 / 청결
- 가정 연계 / 가정연계 / 부모연계 (학부모 소통용 질문이라 활동 아님)

포함해야 할 카테고리 (실제 활동):
- 놀이 / 자유놀이 / 오전자유놀이 / 오후자유놀이
- 바깥놀이 / 실외놀이 / 실내대체
- 기본생활습관 (예: "줄을 서서 기다려요")
- 안전교육 (예: "폭염 비상대응", "(성폭력)내가 좋아하는 것들")
- 특별활동 / 오감놀이 / 체육 / 영어

카테고리 판단은 표의 첫 번째 열(또는 첫 행) 카테고리 이름을 보고 결정. 카테고리 이름이 위 제외 목록에 매치되면 그 행/열의 셀은 모두 스킵.

【3단계: 그 외 정보】
- year, month: 서두에서 (예: "2026년 8월"). 못 찾으면 null.
- theme: 그 달의 큰 놀이 주제 (예: "동물이랑 즐겁게 놀아요"). 서두 또는 표 상단에.
- subtheme: 그 주의 소주제 (표에 "예상 놀이" 행/열 등). 없으면 빈 문자열.

【절대 금지】
- **주(week) 배정을 활동 내용의 의미로 추측하지 마.** 반드시 표의 위치(열 번호 또는 행 번호)로만 판단.
- 3주 활동을 2주에 넣거나, 4주 활동을 3주에 넣는 등 위치 오배정.
- 서두 텍스트의 나열만으로 활동을 뽑아 여러 주에 뿌리는 것. 반드시 표에서 뽑아.
- 카테고리 이름("놀이", "바깥놀이" 등)을 activities 배열에 넣는 것.
- 표 방향(A vs B)을 문서 안에서 왔다갔다 하는 것. 한 표에서는 한 방향만 적용.
- **활동 순서를 임의로 바꾸는 것.** 원본 문서 순서 그대로 배열에 넣어.

【출력】
반드시 아래 형식의 순수 JSON으로만 응답 (마크다운 코드블록/설명 문장 금지):
{{
  "year": 2026,
  "month": 8,
  "theme": "동물이랑 즐겁게 놀아요",
  "weeks": [
    {{
      "weekNumber": 1,
      "subtheme": "다양한 동물이 궁금해요",
      "activities": ["동물 모양 판에 공을 던져요", "동물 이름에 끼적여요", "(바깥놀이)자연물로 동물을 만들어요"]
    }},
    {{
      "weekNumber": 2,
      "subtheme": "동물 가족이 궁금해요",
      "activities": ["엄마 캥거루가 되어요", "..."]
    }}
  ]
}}

문서 텍스트:
{raw_text}
"""

    instructions = (
        "너는 어린이집 월간 놀이 계획안을 구조화된 JSON으로 정리하는 파서야. "
        "먼저 표의 주(week) 배치 방향(열=주 vs 행=주)을 판단하고, "
        "그 다음 표의 셀 위치(열 번호 또는 행 번호)로만 주를 배정해. "
        "활동 내용의 의미로 주를 추측하는 것은 절대 금지. "
        "반드시 순수 JSON만 응답하고, 마크다운 코드블록이나 설명 문장은 절대 붙이지 마. "
        "문서에 없는 정보는 임의로 만들지 말고 null 또는 빈 배열로 남겨."
    )

    try:
        response = client_observation.responses.create(
            model=model_name,
            instructions=instructions,
            input=prompt,
            max_output_tokens=3000,
        )
        raw = _extract_json(response.output_text)
        parsed = json.loads(raw)
    except json.JSONDecodeError as e:
        print("MONTHLY PLAN JSON ERROR:", repr(e))
        raise HTTPException(status_code=500, detail="AI 응답을 해석하지 못했어요.")
    except Exception as e:
        print("MONTHLY PLAN ERROR:", repr(e))
        raise HTTPException(status_code=500, detail=f"월간계획안 파싱 실패: {repr(e)}")

    # 정규화: 필수 키 보장
    weeks = parsed.get("weeks") or []
    normalized_weeks = []
    for w in weeks:
        if not isinstance(w, dict):
            continue
        acts = w.get("activities") or []
        acts = [str(a).strip() for a in acts if str(a).strip()]
        normalized_weeks.append({
            "weekNumber": int(w.get("weekNumber") or (len(normalized_weeks) + 1)),
            "subtheme": str(w.get("subtheme") or "").strip(),
            "activities": acts,
        })

    return {
        "year": parsed.get("year"),
        "month": parsed.get("month"),
        "theme": str(parsed.get("theme") or "").strip(),
        "weeks": normalized_weeks,
    }


# ──────────────── 주간보육일지 DOCX 내보내기 ────────────────

class WeeklyDiaryDaySlots(BaseModel):
    morningFree: str = ""
    outdoor: str = ""
    special: str = ""


class ExportWeeklyDiaryRequest(BaseModel):
    class_name: str = ""
    year: int
    month: int
    week_number: int
    age_group: str = ""
    teacher_name: str = ""
    director_name: str = ""
    theme: str = ""
    subtheme: str = ""
    expectations: str = ""
    days: dict  # { "월": {"morningFree": "...", ...}, ... }
    evaluations: dict  # { "월": "...", ... }
    dates: dict  # { "월": "2026-08-03", ... }


_WEEK_ORDER = ["월", "화", "수", "목", "금", "토"]

# 슬롯 정의: (label, type, standard_text_or_None, extra_rows)
# type: 'text' = 병합 표준 문구
#       'morningFree' = 앱 데이터의 morningFree 를 병합 셀에 불릿 리스트로
#       'outdoor'     = outdoor 병합 셀 + 별도 (O) 행 (요일별)
#       'special'     = 요일별 활동명 셀 + 별도 (O) 행 (요일별)
# extra_rows: 라벨 아래 추가로 필요한 데이터 행 수 (0 = 라벨과 동일 행에만 컨텐츠)
_SLOT_DEFS = [
    ("등원 및 통합보육\n(07:30~09:00)",
     "text", "어린이집에 오면서 보았던 이야기를 나눠요", 0),
    ("손씻기 및 오전 간식\n(09:00~09:30)",
     "text", "<오전 간식 식단표 참고>\n- 손씻기 후 순차 제공", 0),
    ("오 전\n자 유\n놀 이\n(09:30~\n10:30)",
     "morningFree", None, 0),
    ("실외놀이 및\n대체활동\n(10:30~11:40)",
     "outdoor", None, 1),   # +1: (O) 마크 별도 행
    ("점심식사 손씻고\n이닦기\n기저귀갈이 및 배변\n(11:40~12:10)",
     "text", "<점심 식단표 참고>\n- 스스로 손씻기와 이닦기", 0),
    ("특별활동\n(12:10~12:40)",
     "special", None, 1),   # +1: (O) 마크 별도 행
    ("낮잠 및 휴식\n(12:40~14:30)",
     "text", "바르게 누워 편하게 쉬어요\n(화장실 다녀오기 / 세안 / 동화책 듣고 낮잠)", 0),
    ("손씻기 및 오후 간식\n(14:30~15:00)",
     "text", "<오후 식단표 참고>\n- 화장실 가기 및 손씻기", 0),
    ("오후 자유놀이\n(15:00~17:00)",
     "text", "오전에 진행한 활동을 연계하여 자유 놀이를 진행한다.", 0),
    ("통합보육 및 귀가\n(17:00~19:30)",
     "text", "귀가 인사 및 부모님과의 연계", 0),
]


def _saturday_has_content(days: dict, evaluations: dict) -> bool:
    """토요일에 활동 슬롯 or 총평 중 하나라도 내용이 있으면 True."""
    if isinstance(days, dict):
        sat = days.get("토")
        if isinstance(sat, dict):
            for k in ("morningFree", "outdoor", "special"):
                if str(sat.get(k, "") or "").strip():
                    return True
    if isinstance(evaluations, dict):
        if str(evaluations.get("토", "") or "").strip():
            return True
    return False


def build_weekly_diary_docx(req: ExportWeeklyDiaryRequest) -> bytes:
    """
    원본 PDF 레이아웃(A4 세로, 라벨+6일 컬럼)을 그대로 재현.
    - 하나의 큰 표에 제목/정보/활동을 모두 배치 (PDF와 동일)
    - 토요일 비활성 시 토 컬럼 전체를 세로 병합해 "등원하는 영아가 없음" 한 번만 표시
    - 페이지 2: 총평 표 + 투약/출석/청결
    """
    sat_active = _saturday_has_content(req.days, req.evaluations)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    n_days = len(_WEEK_ORDER)  # 6
    n_cols = 1 + n_days         # 라벨 + 월~토 = 7

    # ── 컬럼 폭 (PDF 실측 기반) ──
    # PDF: 라벨 94pt, 요일 각 79pt, 마지막 토 52pt (총 566pt)
    # A4 세로 여백 제외 폭 ≈ 19cm → 라벨 3.2 / 월~금 각 2.65 / 토 1.75
    label_w = Cm(3.2)
    day_w = Cm(2.65)
    sat_w = Cm(1.75)
    def _apply_widths(row):
        row.cells[0].width = label_w
        for i in range(1, n_cols - 1):
            row.cells[i].width = day_w
        row.cells[n_cols - 1].width = sat_w

    # ── 대형 통합 표: 제목 + 정보 3행 + 요일헤더 + 슬롯 (가변 행) ──
    activity_rows = sum(1 + s[3] for s in _SLOT_DEFS)  # 각 슬롯 1행 + 추가행
    total_rows = 1 + 3 + 1 + activity_rows  # title + info×3 + dayHeader + slots
    tbl = doc.add_table(rows=total_rows, cols=n_cols)
    tbl.style = "Table Grid"
    for r in tbl.rows:
        _apply_widths(r)

    # (0) 제목 (전체 병합)
    title_row = tbl.rows[0]
    m = title_row.cells[0]
    for i in range(1, n_cols):
        m = m.merge(title_row.cells[i])
    title_text = f"<{req.year}년 {req.month}월 {req.week_number}주> {req.class_name or '반'} 보  육  일  지"
    if req.age_group:
        title_text += f"  ({req.age_group})"
    _set_cell(m, title_text, bold=True, center=True, size=13)
    _set_row_min_height(title_row, Cm(1.1))

    # (1~3) 정보 행: 놀이 주제 / 예상 놀이 / 교사 기대
    def _info_row(r_idx: int, label: str, value: str, min_h: float = 0.9):
        row = tbl.rows[r_idx]
        _set_cell(row.cells[0], label, bold=True, center=True, size=10)
        m = row.cells[1]
        for i in range(2, n_cols):
            m = m.merge(row.cells[i])
        _set_cell(m, value or "", size=10)
        _set_row_min_height(row, Cm(min_h))

    _info_row(1, "놀이 주제", req.theme)
    _info_row(2, "예상 놀이", req.subtheme)
    _info_row(3, "교사의 기대", req.expectations, min_h=1.1)

    # (4) 요일 헤더 행
    day_header_idx = 4
    day_row = tbl.rows[day_header_idx]
    _set_cell(day_row.cells[0], "요일\n활 동", bold=True, center=True, size=10)
    for i, day in enumerate(_WEEK_ORDER):
        date = req.dates.get(day, "") if isinstance(req.dates, dict) else ""
        header = f"{date[5:]}({day})" if date else day
        _set_cell(day_row.cells[1 + i], header, bold=True, center=True, size=9)
    _set_row_min_height(day_row, Cm(0.9))

    # ── 슬롯 행들 (가변 행 수) ──
    # 각 슬롯의 모든 행 인덱스 저장 → 나중에 토 컬럼 세로 병합에 사용
    all_slot_row_indices = []  # 모든 슬롯의 행 인덱스 (평탄화)
    cur_row = day_header_idx + 1

    end_col = n_cols if sat_active else n_cols - 1  # 병합 종점(포함하지 않음)
    day_count = n_days if sat_active else n_days - 1

    def _merge_content(row, text: str, size: int = 9):
        m = row.cells[1]
        for i in range(2, end_col):
            m = m.merge(row.cells[i])
        _set_cell(m, text or "", size=size)

    for label, slot_type, standard, extra in _SLOT_DEFS:
        slot_rows_idx = list(range(cur_row, cur_row + 1 + extra))
        all_slot_row_indices.extend(slot_rows_idx)

        # 라벨 셀: 슬롯의 여러 행을 세로 병합
        top_row = tbl.rows[slot_rows_idx[0]]
        label_cell = top_row.cells[0]
        for r_idx in slot_rows_idx[1:]:
            label_cell = label_cell.merge(tbl.rows[r_idx].cells[0])
        _set_cell(label_cell, label, bold=True, center=True, size=9)

        if slot_type == "text":
            _merge_content(top_row, standard or "")
            _set_row_min_height(top_row, Cm(1.1))
        elif slot_type == "morningFree":
            all_activities = _collect_unique(req.days, "morningFree")
            _merge_content(top_row, _as_bullets(all_activities) or "-")
            _set_row_min_height(top_row, Cm(4.5))
        elif slot_type == "outdoor":
            all_outdoor = _collect_unique(req.days, "outdoor")
            _merge_content(top_row, _as_bullets(all_outdoor) or "-")
            _set_row_min_height(top_row, Cm(2.2))
            # 별도 (O) 마크 행
            mark_row = tbl.rows[slot_rows_idx[1]]
            for i in range(day_count):
                day = _WEEK_ORDER[i]
                mark = "(O)" if _get_slot(req.days, day, "outdoor").strip() else ""
                _set_cell(mark_row.cells[1 + i], mark, center=True, size=9)
            _set_row_min_height(mark_row, Cm(0.6))
        elif slot_type == "special":
            # 활동명 행 (요일별)
            for i in range(day_count):
                day = _WEEK_ORDER[i]
                content = _get_slot(req.days, day, "special").strip()
                _set_cell(top_row.cells[1 + i], content, center=True, size=9)
            _set_row_min_height(top_row, Cm(0.9))
            # (O) 마크 행
            mark_row = tbl.rows[slot_rows_idx[1]]
            for i in range(day_count):
                day = _WEEK_ORDER[i]
                mark = "(O)" if _get_slot(req.days, day, "special").strip() else ""
                _set_cell(mark_row.cells[1 + i], mark, center=True, size=9)
            _set_row_min_height(mark_row, Cm(0.6))

        cur_row += 1 + extra

    # ── 토요일 비활성일 때: 토 컬럼 전체를 세로 병합해 한 셀로 ──
    if not sat_active:
        sat_col_idx = n_cols - 1
        first_sat_cell = tbl.rows[all_slot_row_indices[0]].cells[sat_col_idx]
        for r_idx in all_slot_row_indices[1:]:
            first_sat_cell = first_sat_cell.merge(tbl.rows[r_idx].cells[sat_col_idx])
        _set_cell(first_sat_cell, "등\n원\n하\n는\n\n영\n아\n가\n\n없\n음",
                  bold=True, center=True, size=10)

    # 담임/원장은 표 상단 우측에
    _insert_top_signatures(doc, req.teacher_name, req.director_name)

    # ── 페이지 나누기 ──
    doc.add_page_break()

    # ── 총평 + 투약일지 + 출석 + 실내청결 통합 표 ──
    eval_days = ["월", "화", "수", "목", "금"]
    if sat_active:
        eval_days.append("토")
    bottom_days = n_days if sat_active else n_days - 1  # 5 or 6
    bottom_cols = 1 + bottom_days  # 라벨 + 요일

    # 통합 표: eval 행수 + 투약/출석/실내청결 3행
    combined_rows = len(eval_days) + 3
    combined = doc.add_table(rows=combined_rows, cols=bottom_cols)
    combined.style = "Table Grid"

    # 컬럼 폭 통일
    label_w2 = Cm(2.5)
    day_w2 = Cm((21 - 2.0 - 2.5) / bottom_days)
    for row in combined.rows:
        row.cells[0].width = label_w2
        for i in range(1, bottom_cols):
            row.cells[i].width = day_w2

    # (A) eval 행들: 첫 열 라벨 병합 세로, 둘째 열 요일 letter, 나머지 병합해 서술
    for i, day in enumerate(eval_days):
        row = combined.rows[i]
        # 둘째 열: 요일
        _set_cell(row.cells[1], day, bold=True, center=True, size=10)
        # 셋째 이후 병합해 서술
        m = row.cells[2]
        for j in range(3, bottom_cols):
            m = m.merge(row.cells[j])
        text = str(req.evaluations.get(day, "") or "").strip() if isinstance(req.evaluations, dict) else ""
        _set_cell(m, text, size=9)
        _set_row_min_height(row, Cm(3.0))

    # 첫 열 (총평 라벨) 세로 병합
    eval_label_cell = combined.rows[0].cells[0]
    for i in range(1, len(eval_days)):
        eval_label_cell = eval_label_cell.merge(combined.rows[i].cells[0])
    _set_cell(eval_label_cell, "총평 및\n활동평가\n특이사항 등",
              bold=True, center=True, size=10)

    # (B) 투약일지 행
    med_idx = len(eval_days)
    r0 = combined.rows[med_idx]
    _set_cell(r0.cells[0], "투 약 일 지", bold=True, center=True, size=10)
    for i in range(bottom_days):
        _set_cell(r0.cells[1 + i], "", center=True, size=9)
    _set_row_min_height(r0, Cm(1.4))

    # (C) 출석(결석) 행
    att_idx = med_idx + 1
    r1 = combined.rows[att_idx]
    _set_cell(r1.cells[0], "출석(결석)", bold=True, center=True, size=10)
    for i in range(bottom_days):
        _set_cell(r1.cells[1 + i], "", center=True, size=9)
    _set_row_min_height(r1, Cm(0.9))

    # (D) 실 내 청 결 행
    clean_idx = att_idx + 1
    r2 = combined.rows[clean_idx]
    _set_cell(r2.cells[0], "실 내 청 결", bold=True, center=True, size=10)
    m = r2.cells[1]
    for i in range(2, bottom_cols):
        m = m.merge(r2.cells[i])
    _set_cell(m, "*교실 / 화장실 청소\n*자체소독 – 교실 / 교구장 및 놀잇감\n*침구(가정으로 보내기)", size=9)
    _set_row_min_height(r2, Cm(1.4))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _collect_unique(days: dict, key: str) -> list:
    """모든 요일의 slot 값들을 라인 단위로 모아 중복 제거."""
    seen: set = set()
    out: list = []
    if not isinstance(days, dict):
        return out
    for day in _WEEK_ORDER:
        d = days.get(day)
        if not isinstance(d, dict):
            continue
        text = str(d.get(key, "") or "")
        for line in text.split("\n"):
            line = line.strip().lstrip("- ").strip()
            if line and line not in seen:
                seen.add(line)
                out.append(line)
    return out


def _as_bullets(items: list) -> str:
    if not items:
        return ""
    return "\n".join(f"- {it}" for it in items)


def _insert_top_signatures(doc: Document, teacher: str, director: str):
    """
    문서의 맨 앞(첫 요소보다 위)에 담임/원장 서명란을 우측 정렬 문단으로 추가.
    (Document.paragraphs 조작으로 삽입)
    """
    from docx.oxml.ns import qn as _qn
    body = doc.element.body
    para = doc.paragraphs[0] if doc.paragraphs else None
    new_p_xml = etree.SubElement(body, _qn("w:p"))
    # 우측 정렬
    pPr = etree.SubElement(new_p_xml, _qn("w:pPr"))
    jc = etree.SubElement(pPr, _qn("w:jc"))
    jc.set(_qn("w:val"), "right")
    run = etree.SubElement(new_p_xml, _qn("w:r"))
    rPr = etree.SubElement(run, _qn("w:rPr"))
    sz = etree.SubElement(rPr, _qn("w:sz"))
    sz.set(_qn("w:val"), "20")  # half-points
    rFonts = etree.SubElement(rPr, _qn("w:rFonts"))
    rFonts.set(_qn("w:eastAsia"), "맑은 고딕")
    t = etree.SubElement(run, _qn("w:t"))
    t.text = f"담임: {teacher or ''}      원장: {director or ''}"
    t.set(_qn("xml:space"), "preserve")
    # 방금 추가한 새 문단을 첫 요소 앞으로 이동
    if para is not None:
        body.remove(new_p_xml)
        para._element.addprevious(new_p_xml)


def _get_slot(days: dict, day: str, key: str) -> str:
    if not isinstance(days, dict):
        return ""
    d = days.get(day)
    if not isinstance(d, dict):
        return ""
    return str(d.get(key, "") or "")


def _force_east_asia_font(run, font_name: str):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), font_name)


@app.post("/export-weekly-diary-docx")
def export_weekly_diary_docx(req: ExportWeeklyDiaryRequest):
    try:
        docx_bytes = build_weekly_diary_docx(req)
        filename = f"주간보육일지_{req.year}년{req.month}월{req.week_number}주_{req.class_name or '반'}.docx"
        encoded_filename = quote(filename, safe="")
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
        )
    except Exception as error:
        print("WEEKLY DIARY DOCX ERROR:", repr(error))
        raise HTTPException(status_code=500, detail=f"docx 생성 중 오류: {repr(error)}")


# ──────────────── 주간보육일지 총평 생성 ────────────────

class WeeklyDayContext(BaseModel):
    day: str  # "월", "화", ...
    date: str  # YYYY-MM-DD
    morning_free: str = ""
    outdoor: str = ""
    special: str = ""
    common_records: List[str] = []      # 반 공통 알림장 본문들
    masked_personal_records: List[str] = []  # 개인 알림장 (이름 마스킹된 본문들)


class GenerateWeeklyEvaluationsRequest(BaseModel):
    class_name: str = ""
    year: int
    month: int
    week_number: int
    age_group: str = ""
    theme: str = ""
    subtheme: str = ""
    style_guide: str = ""
    days: List[WeeklyDayContext]


def _resolve_age_context(age_group: str) -> tuple[str, str]:
    """(표기용 연령, 아이 지칭어) — 만 3세 이상은 '유아', 그 미만은 '영아'."""
    raw = (age_group or "").strip()
    display = raw or "만 2세"
    m = re.search(r"(\d+)", raw)
    if m:
        age_num = int(m.group(1))
        return display, ("유아" if age_num >= 3 else "영아")
    return display, "영아"


@app.post("/generate-weekly-evaluations")
def generate_weekly_evaluations(request: GenerateWeeklyEvaluationsRequest):
    if not api_key:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY가 설정되지 않았습니다.")

    if not request.days:
        raise HTTPException(status_code=400, detail="요일 정보가 비어 있습니다.")

    # 요일별 컨텍스트 구성
    day_blocks: List[str] = []
    for d in request.days:
        activities = []
        if d.morning_free.strip():
            activities.append(f"[오전 자유놀이] {d.morning_free.strip()}")
        if d.outdoor.strip():
            activities.append(f"[실외놀이] {d.outdoor.strip()}")
        if d.special.strip():
            activities.append(f"[특별활동] {d.special.strip()}")
        activities_text = "\n".join(activities) if activities else "(등록된 활동 없음)"

        record_blocks = []
        for r in d.common_records:
            if r.strip():
                record_blocks.append(f"- (반 공통) {r.strip()}")
        for r in d.masked_personal_records:
            if r.strip():
                record_blocks.append(f"- (개인, 이름 마스킹) {r.strip()}")
        if record_blocks:
            records_text = (
                "⭐ 이 요일에 실제로 저장된 알림장입니다. "
                "총평은 반드시 아래 알림장에 나온 활동·아이의 반응·표현·교사의 상호작용을 근거로 작성하세요. "
                "새로운 활동을 지어내지 말고 알림장 내용을 우선하세요.\n"
                + "\n".join(record_blocks)
            )
        else:
            records_text = "(해당 요일에 저장된 알림장 없음 — 활동명과 주제만으로 자연스러운 관찰문을 구성)"

        day_blocks.append(
            f"### {d.day}요일 ({d.date})\n"
            f"활동:\n{activities_text}\n"
            f"참고 알림장:\n{records_text}"
        )

    context_text = "\n\n".join(day_blocks)

    age_display, child_word = _resolve_age_context(request.age_group)

    prompt = f"""
{age_display} {child_word}의 어린이집 주간보육일지 '실행기록 및 평가'를 요일별로 작성해줘.

【상황】
- 반: {request.class_name or '지정 안됨'}
- 대상 연령: {age_display}
- {request.year}년 {request.month}월 {request.week_number}주
- 놀이 주제: {request.theme or '(없음)'}
- 예상 놀이: {request.subtheme or '(없음)'}

【요일별 정보】
{context_text}

【작성 기준】
- 어린이집 보육일지에 바로 사용할 수 있는 문체로 작성.
- 활동의 실제 진행 모습 → {child_word}의 관심과 반응 → 교사의 상호작용 및 놀이지원 → 놀이의 확장 → 활동을 통해 경험한 내용 순으로 자연스럽게 한 문단으로 서술.
- {child_word}가 실제 놀이하는 장면이 떠오르도록 구체적으로 작성.
- {child_word}의 행동을 "관심을 보였다", "탐색하는 모습을 보였다", "반복하여 시도하였다", "즐겁게 참여하였다" 등 관찰 중심 표현으로 기록.
- 교사의 일방적인 설명보다 {child_word}의 자발적인 탐색과 놀이 과정을 중심으로 작성.
- 필요한 부분에 (놀이지원)을 넣고 교사가 어떻게 상호작용하거나 놀이를 확장했는지 기록.
- {child_word}의 언어 수준을 고려해 짧은 단어, 의성어·의태어, 간단한 문장으로 표현하는 모습을 자연스럽게 포함.
- 신체·언어·감각·탐색·사회관계·예술경험 등 활동과 관련된 발달 경험이 자연스럽게 드러나도록 작성.
- 지나치게 교육적이거나 초등학생 수준의 표현은 피하고 {age_display} 발달 수준에 맞게 작성.
- "학습하였다", "정확하게 이해하였다" 대신 "관심을 가졌다", "경험하였다", "탐색하였다", "표현하였다" 사용.
- 친구와 함께하는 활동이면 또래를 관찰하거나 따라 하기, 함께 놀이하기, 차례 경험하기 등의 모습을 자연스럽게 포함.
- 미술활동은 재료 탐색 → 자유로운 표현 → 교사의 지원 → 완성 후 반응이 드러나도록 작성.
- 신체활동은 신체 움직임, 대근육 사용, 균형 및 움직임 조절 등이 자연스럽게 드러나도록 작성.
- 끼적이기 활동은 그림·사진 탐색 → 크레파스 등의 도구 선택 → 자유로운 끼적이기 → 관련 사물이나 경험에 대한 언어적 상호작용 순으로 작성.
- 바깥놀이는 자연환경이나 놀이기구 탐색, 신체활동, 안전한 놀이를 위한 교사의 지원이 드러나도록 작성.
- 마지막에는 활동을 통해 {child_word}가 경험한 내용을 자연스럽게 평가하며 마무리.
- **분량은 각 요일 350~500자.**
- 별도의 목표·준비물·활동방법 등은 작성하지 말고 '실행기록 및 평가' 내용만.
- "참고 알림장"이 있으면 그 내용을 반영해 실제 있었던 관찰을 자연스럽게 서술.
- 참고 알림장이 없으면 활동 이름과 계획된 주제만 기반으로 자연스러운 관찰문을 상상해 서술(구체 사실은 만들지 말고 활동 성격에 맞는 일반적 관찰로).
- 개인 알림장은 이름이 [영아]로 마스킹돼 있으니 특정 아이를 지칭하지 말고 반 전체의 관찰로 서술.
- 활동이 아예 없는 요일이면 빈 문자열로 남겨.

【참고 예시 — 활동명 "호랑이에 대해 알고 호랑이를 꾸며요" (만 2세 영아 기준)】
호랑이 사진을 제공하자 영아는 호랑이의 얼굴과 몸에 있는 줄무늬를 살펴보며 관심을 보였다. 교사는 "호랑이는 어떤 색일까?", "몸에는 어떤 무늬가 있니?"라고 질문하며 호랑이의 생김새와 특징을 탐색할 수 있도록 상호작용하였다.(놀이지원) 영아는 "어흥" 하며 호랑이 울음소리를 흉내 내거나 손과 몸을 움직여 호랑이의 모습을 표현하며 즐거워하였다. 이후 호랑이 그림과 다양한 꾸미기 재료를 제공하자 영아는 크레파스로 줄무늬를 그리거나 색종이와 스티커를 붙이며 자유롭게 호랑이를 꾸몄다. 교사는 영아의 표현을 긍정적으로 격려하고 호랑이의 생김새와 움직임에 대해 이야기하며 놀이를 확장하였다. 활동을 통해 호랑이의 특징에 관심을 가지고 다양한 재료를 활용하여 창의적으로 표현하는 경험을 할 수 있었다.
(위 예시는 문체·구성·분량 참고용. 실제 대상 연령이 {age_display}이면 표현 수준을 그에 맞게 조정하고, {child_word} 지칭어도 그대로 사용.)

{f'【문체 가이드】{chr(10)}{request.style_guide.strip()}' if request.style_guide.strip() else ''}

【출력】
반드시 아래 형식의 순수 JSON으로만 응답 (마크다운 코드블록/설명 금지):
{{
  "월": "월요일 실행기록 및 평가...",
  "화": "화요일 실행기록 및 평가...",
  "수": "...",
  "목": "...",
  "금": "...",
  "토": "..."
}}

응답에 요일 키는 위 요일별 정보에 등장한 요일만 포함해.
"""

    instructions = (
        "너는 어린이집 주간보육일지 총평을 쓰는 어시스턴트야. "
        "실제 관찰에 없는 구체 사실(특정 아이 이름/발화 원문/부상 사고 등)은 만들어내지 마. "
        "반드시 순수 JSON만 응답하고, 마크다운 코드블록이나 설명 문장은 절대 붙이지 마."
    )

    try:
        response = client_observation.responses.create(
            model=model_name,
            instructions=instructions,
            input=prompt,
            max_output_tokens=3500,
        )
        raw = _extract_json(response.output_text)
        parsed = json.loads(raw)
        if not isinstance(parsed, dict):
            raise ValueError("응답이 객체가 아님")
        return {k: str(v).strip() for k, v in parsed.items() if isinstance(k, str)}
    except json.JSONDecodeError as e:
        print("WEEKLY EVAL JSON ERROR:", repr(e))
        raise HTTPException(status_code=500, detail="AI 응답을 해석하지 못했어요.")
    except Exception as e:
        print("WEEKLY EVAL ERROR:", repr(e))
        raise HTTPException(status_code=500, detail=f"총평 생성 실패: {repr(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
